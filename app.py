import openpyxl  # Required for saving Excel files
from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
import cv2
import numpy as np
import easyocr
from PIL import Image
import io
import re
import time
import arabic_reshaper
from bidi.algorithm import get_display  # Required for Arabic text formatting
import multiprocessing
from docx import Document
from werkzeug.utils import secure_filename
import pandas as pd
import os
from pdf2image import convert_from_bytes
from flask import send_file
import mimetypes
import tempfile
import threading
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from paddleocr import PaddleOCR
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor

# Executor for parallel execution
executor = concurrent.futures.ThreadPoolExecutor(max_workers=4)



app = Flask(__name__)
CORS(app)  # Enable CORS for frontend communication

# Print all registered routes
with app.test_request_context():
    print(app.url_map)

@app.route('/')
def home():
    return render_template('index.html')  # Ensure this file exists in "templates" folder

@app.route('/image-to-word')
def image_to_word():
    return render_template('image_to_word.html')  # Ensure this template exists

@app.route('/image-to-excel')
def image_to_excel():
    return render_template('image_to_excel.html')  # Ensure this template exists

@app.route('/pdf-to-image')
def pdf_to_image():
    return render_template('pdf_to_image.html')  # Ensure this template exists
    
# ✅ Define supported languages (Tesseract & EasyOCR)
supported_languages = [
    "en", "hin", "spa", "fra", "deu", "ara", "por", "rus", "jpn", "ita", "kor", "tur", "nld",
    "pol", "vie", "urd", "ben", "tam", "tel", "ch_sim", "ch_tra"  # Add Chinese here
]

# ✅ EasyOCR-supported languages (Filtered)
easyocr_supported_languages = [
    "en", "hi", "es", "fr", "de", "ar", "pt", "ru", "ja", "it", "ko", "tr", "nl",
    "pl", "vi", "ur", "bn", "ta", "te", "ch_sim", "ch_tra"  # Add Chinese here
]


filtered_languages = [lang for lang in supported_languages if lang in easyocr_supported_languages]

# ✅ Initialize EasyOCR Reader
try:
    # Ensure correct language codes
    language_map = {
        "chi_sim": "ch_sim",
        "chi_tra": "ch_tra",
        "jpn": "ja"  # Ensure Japanese uses the correct code
    }

    # Apply mapping
    filtered_languages = [language_map.get(lang, lang) for lang in filtered_languages]

    # Remove duplicates and enforce only one Chinese type
    if "ch_sim" in filtered_languages and "ch_tra" in filtered_languages:
        filtered_languages.remove("ch_sim")  # Prioritize Traditional (ch_tra), or swap if needed

    # Ensure English is included when using Chinese
    if any(lang in filtered_languages for lang in ["ch_sim", "ch_tra"]) and "en" not in filtered_languages:
        filtered_languages.append("en")

    # Ensure Japanese includes English
    if "jpn" in filtered_languages and "en" not in filtered_languages:
        filtered_languages.append("en")

    # Debugging: Print final languages list
    print("✅ Final languages passed to EasyOCR:", filtered_languages)

    # Initialize EasyOCR reader
    reader = easyocr.Reader(filtered_languages)

except ValueError as e:
    print(f"❌ EasyOCR initialization failed: {e}")


# ✅ Function: Clean OCR Output (Better Formatting & Error Correction)
def clean_ocr_text(text):
    text = re.sub(r"\n+", "\n", text)  # Remove extra newlines
    text = re.sub(r"\s+", " ", text)  # Remove extra spaces
    text = re.sub(r"[^\w\s,.!?؀-ۿ]", "", text, flags=re.UNICODE)  # Keep Urdu & other characters
    text = text.strip()
    return text

def clean_japanese_text(text):
    text = re.sub(r'\s+', '', text)  # Remove extra spaces
    return text.strip()


# ✅ Function: CLAHE Contrast Enhancement (Only for Arabic OCR)
def apply_clahe(image):
    lab = cv2.cvtColor(image, cv2.COLOR_BGR2LAB)  # Convert to LAB color space
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))  # Apply CLAHE
    cl = clahe.apply(l)
    merged = cv2.merge((cl, a, b))
    return cv2.cvtColor(merged, cv2.COLOR_LAB2BGR)  # Convert back to BGR

# ✅ Function: Adjust Contrast for Better OCR
def adjust_contrast(image, language):
    alpha = 1.5  # Contrast control (1.0-3.0)
    beta = 10    # Brightness control (0-100)
    adjusted = cv2.convertScaleAbs(image, alpha=alpha, beta=beta)

    if language == "ara":  # Apply CLAHE only for Arabic
        adjusted = apply_clahe(adjusted)

    return adjusted

# ✅ Function: Noise Removal (For Smoother Characters)
def remove_noise(image):
    return cv2.fastNlMeansDenoising(image, None, 30, 7, 21)

# ✅ Function: Enhance Text Size (For Better Recognition)
def enlarge_text(image):
    scale_percent = 150  # Enlarge by 150%
    width = int(image.shape[1] * scale_percent / 100)
    height = int(image.shape[0] * scale_percent / 100)
    resized = cv2.resize(image, (width, height), interpolation=cv2.INTER_CUBIC)
    return resized

# ✅ Function: Line Detection (Removes Unwanted Lines)
def detect_lines(image):
    edges = cv2.Canny(image, 50, 150, apertureSize=3)
    lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=100, minLineLength=30, maxLineGap=5)
    if lines is not None:
        for line in lines:
            x1, y1, x2, y2 = line[0]
            cv2.line(image, (x1, y1), (x2, y2), (255, 255, 255), 2)
    return image

# ✅ Function: Correct Skew (Disabled for Urdu)
def correct_skew(image, language):
    if language == "urd":
        return image  # Don't apply skew correction for Urdu

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    coords = np.column_stack(np.where(gray > 0))
    angle = cv2.minAreaRect(coords)[-1]

    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle

    (h, w) = image.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
    return rotated

# ✅ Function: Preprocess Image for OCR
def preprocess_image(image_cv, language):
    gray = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)  # Convert to grayscale
    adjusted = adjust_contrast(gray, language)  # Improve contrast
    rotated = correct_skew(adjusted, language)  # Fix skew (except Urdu)
    denoised = remove_noise(rotated)  # Remove noise
    enlarged = enlarge_text(denoised)  # Enlarge text
    lined_image = detect_lines(enlarged)  # Detect & remove unwanted lines

    blurred = cv2.GaussianBlur(lined_image, (5, 5), 0)  # Reduce noise

    # ✅ Different Thresholding for Urdu
    if language == "urd":
        _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    else:
        thresh = cv2.adaptiveThreshold(
            blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )

    kernel = np.ones((1, 1), np.uint8)
    morph = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)

    return morph

def preprocess_chinese(image_cv):
    """Preprocess image for better OCR on Chinese text (Simplified & Traditional)"""
    
    # Convert to grayscale
    gray = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)

    # Improve contrast using Adaptive Histogram Equalization
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)

    # Apply Gaussian Blur to remove noise
    blurred = cv2.GaussianBlur(enhanced, (3, 3), 0)

    # Sharpen the image to enhance thin strokes of Chinese characters
    kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
    sharpened = cv2.filter2D(blurred, -1, kernel)

    # Apply Adaptive Thresholding (better for Chinese characters)
    binary = cv2.adaptiveThreshold(
        sharpened, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, 5
    )

    # Morphological Transformations to clean noise (suitable for thin strokes)
    kernel = np.ones((2, 2), np.uint8)
    morph = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

    return morph



# ✅ Function: Arabic-Specific Preprocessing (Better Accuracy)
def preprocess_arabic(image_cv):
    gray = cv2.cvtColor(image_cv, cv2.COLOR_BGR2GRAY)  # Convert to grayscale
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)  # Reduce noise
    _, binary = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)  # Binarization

    # Morphological Transformations to connect Arabic letters properly
    kernel = np.ones((2, 2), np.uint8)
    morph = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

    return morph

# ✅ Arabic Text Normalization (Fix Grammar & Spacing)
def normalize_arabic_text(text):
    text = re.sub(r"\s+", " ", text)  # Fix extra spaces
    text = re.sub(r"[^\w\s،.!؟؀-ۿ]", "", text, flags=re.UNICODE)  # Keep Arabic characters
    text = arabic_reshaper.reshape(text)  # Fix Arabic letter shaping
    text = get_display(text)  # Correct word order & alignment
    return text.strip()

@app.route("/ocr", methods=["POST"])
def ocr():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    language = request.form.get("language", "en")  # Default to English

    try:
        # Convert Image to OpenCV Format
        image = Image.open(io.BytesIO(file.read()))
        image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

        # ✅ Apply Preprocessing (Special Handling for Arabic, Urdu, Chinese)
        if language in ["ch_sim", "ch_tra"]:
            processed_image = preprocess_chinese(image_cv)  # Optimized preprocessing for Chinese
        elif language in ["ara", "urd"]:
            processed_image = preprocess_arabic(image_cv)
        else:
            processed_image = preprocess_image(image_cv, language)

        # ✅ Configure Tesseract OCR (Optimized for Each Language)
        if language in ["ch_sim", "ch_tra"]:
            tesseract_psm = "7"
            tesseract_config = f"--oem 3 --psm {tesseract_psm} -l {language} -c preserve_interword_spaces=1"
        elif language in ["jpn"]:
            tesseract_psm = "6"
            tesseract_config = "--oem 3 --psm 6 -l jpn"  # Japanese-specific settings
        elif language in ["ara", "urd"]:
            tesseract_psm = "12"
            tesseract_config = f"--oem 3 --psm {tesseract_psm} -l {language}"
        else:
            tesseract_psm = "6"
            tesseract_config = f"--oem 3 --psm {tesseract_psm} -l {language}"

        # ✅ Perform OCR with Tesseract
        try:
            tesseract_text = pytesseract.image_to_string(processed_image, config=tesseract_config).strip()

            # ✅ Debug: Save preprocessed image to check output
            if language in ["ch_sim", "ch_tra"]:
                cv2.imwrite("debug_preprocessed_chinese.png", processed_image)
            else:
                cv2.imwrite("debug_preprocessed.png", processed_image)

        except pytesseract.TesseractError as e:
            print(f"❌ Tesseract Error: {str(e)}")
            tesseract_text = ""

        # ✅ Try PaddleOCR if Tesseract Fails
        if not tesseract_text:
            print(f"⚠️ Warning: Tesseract failed for {language}, trying PaddleOCR...")

            if language in ["ch_sim", "ch_tra"]:
                ocr_engine = PaddleOCR(use_angle_cls=True, lang="ch")  # Optimized for Chinese
                results = ocr_engine.ocr(image_cv, cls=True)

                # Extract text from PaddleOCR results
                paddle_text = " ".join([line[1][0] for result in results for line in result])

                if paddle_text.strip():
                    tesseract_text = paddle_text  # Use PaddleOCR result if successful
                else:
                    print("⚠️ PaddleOCR also failed. Trying EasyOCR...")

        # ✅ Try EasyOCR if both Tesseract and PaddleOCR fail
        easyocr_text = ""
        if not tesseract_text:
            if language in ["ch_sim", "ch_tra"]:
                reader = easyocr.Reader(["ch_tra"])  # Use Traditional Chinese model
                easyocr_text = "".join(reader.readtext(image_cv, detail=0))

                # Final cleanup of OCR text
                easyocr_text = clean_ocr_text(easyocr_text)
            elif language in ["jpn", "ara", "urd"]:
                if language == "jpn":
                    reader = easyocr.Reader(["ja"])  # Japanese model
                    easyocr_text = "".join(reader.readtext(image_cv, detail=0))
                    easyocr_text = clean_japanese_text(easyocr_text)  # Clean Japanese text
                else:
                    reader = easyocr.Reader([language])
                    easyocr_text = " ".join(reader.readtext(image_cv, detail=0))

                    if language in ["ara", "urd"]:
                        easyocr_text = normalize_arabic_text(easyocr_text)
                    else:
                        easyocr_text = clean_ocr_text(easyocr_text)
            else:
                easyocr_text = "❌ EasyOCR does not support this language."

        # ✅ Return Best Result
        return jsonify({
            "tesseract": clean_ocr_text(tesseract_text),
            "easyocr": easyocr_text.strip() if easyocr_text else "✅ EasyOCR was not required."
        })

    except IOError:
        return jsonify({"error": "File reading error. Please upload a valid image."}), 400
    except ValueError as e:
        return jsonify({"error": f"Processing error: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500


# Define upload & output directories
UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
OUTPUT_DIR = "downloads"

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER  

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Language code mapping (your code -> EasyOCR code)
LANGUAGE_CODE_MAPPING = {
    "en": "en",  # English
    "rus": "ru",  # Russian
    "hin": "hi",  # Hindi
    "chi_sim": "ch_sim",  # Simplified Chinese
    "chi_tra": "ch_tra",  # Traditional Chinese
    "ara": "ar",  # Arabic
    "spa": "es",  # Spanish
    "fra": "fr",  # French
    "deu": "de",  # German
    "por": "pt",  # Portuguese
    "jpn": "ja",  # Japanese
    "kor": "ko",  # Korean
    "ita": "it",  # Italian
    "tur": "tr",  # Turkish
    "nld": "nl",  # Dutch
    "pol": "pl",  # Polish
    "vie": "vi",  # Vietnamese
    "urd": "ur",  # Urdu
    "ben": "bn",  # Bengali
    "tam": "ta",  # Tamil
    "tel": "te",  # Telugu
}

def resize_image(image_path, max_width=1000):
    """ Resize image to max width while keeping aspect ratio. """
    try:
        img = Image.open(image_path)
        if img.width > max_width:
            ratio = max_width / img.width
            new_height = int(img.height * ratio)
            img = img.resize((max_width, new_height))
            img.save(image_path)
        return image_path
    except Exception as e:
        print(f"⚠️ Image resizing failed: {e}")
        return image_path  # Return original if resizing fails


def perform_ocr(image_path, selected_language):
    """ Perform OCR using multiprocessing for faster execution. """
    try:
        easyocr_language = LANGUAGE_CODE_MAPPING.get(selected_language, "en")
        print(f"🔍 Using selected language: {easyocr_language} for OCR")

        # Resize the image before OCR (Optional but improves speed)
        image_path = resize_image(image_path)

        # Use ThreadPoolExecutor for parallel processing
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
            reader = easyocr.Reader([easyocr_language], gpu=False, verbose=False)
            future = executor.submit(reader.readtext, image_path, detail=0)
            extracted_text = future.result()  # Wait for OCR to complete

        if not extracted_text:
            print("❌ OCR returned empty text!")
            return None

        print(f"✅ OCR Output: {extracted_text}")
        return "\n".join(extracted_text)

    except Exception as e:
        print(f"❌ Error in OCR processing: {e}")
        return None


from io import BytesIO

def save_text_as_word_fast(text, lang="en"):
    """ Saves extracted OCR text as a Word document in memory. """
    if text is None or text.strip() == "":
        print("❌ No text to save in Word document!")
        return None  

    try:
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run(text)

        # Font selection based on language
        font_name = "Calibri"
        if lang in ["hi", "bn", "ta", "te", "ur"]:
            font_name = "Nirmala UI"
        elif lang in ["es", "fr", "de", "pt", "nl", "pl", "vi"]:
            font_name = "Arial"
        elif lang in ["ru", "ja", "ko", "ch_sim", "ch_tra"]:
            font_name = "Malgun Gothic"

        run.font.name = font_name
        run.font.size = Pt(12)

        # ✅ Save Word file in memory instead of disk
        word_stream = io.BytesIO()
        doc.save(word_stream)
        word_stream.seek(0)

        print("✅ Word file generated in memory")
        return word_stream

    except Exception as e:
        print(f"❌ Error saving Word file: {e}")
        return None

def process_ocr_and_generate_word(file_path, selected_language):
    """Perform OCR and generate Word document asynchronously."""
    try:
        extracted_text = perform_ocr(file_path, selected_language)
        if not extracted_text:
            return None, "OCR Processing Failed!"

        word_buffer = save_text_as_word_fast(extracted_text, selected_language)
        if word_buffer is None:
            return None, "Failed to generate Word document."

        return word_buffer, None  # Return buffer and no error

    except Exception as e:
        return None, str(e)



@app.route("/ocr-to-word", methods=["POST"])
def ocr_to_word():
    """API endpoint for OCR + Word generation."""
    try:
        if "file" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "No file selected"}), 400

        selected_language = request.form.get("language", "en")
        file_path = os.path.join(UPLOAD_FOLDER, secure_filename(file.filename))
        file.save(file_path)

        # Run OCR in a separate process
        future = executor.submit(process_ocr_and_generate_word, file_path, selected_language)
        word_buffer, error = future.result()

        if error:
            return jsonify({"error": error}), 400

        return send_file(word_buffer, as_attachment=True, download_name="extracted_text.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        return jsonify({"error": str(e)}), 500



# ✅ Function to Save Extracted Text as an Excel File
def save_text_as_excel(text, filename):
    excel_path = os.path.join(OUTPUT_DIR, filename)
    df = pd.DataFrame({"Extracted Text": text.split("\n")})  # Convert text into rows
    df.to_excel(excel_path, index=False)  # Save as Excel
    return excel_path


# ✅ Route for Image to Excel OCR
@app.route("/ocr-to-excel", methods=["POST"])
def ocr_to_excel():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    language = request.form.get("language", "eng")
    extracted_text = request.form.get("extracted_text")  # ✅ Fix: Extracted text from frontend

    try:
        if not extracted_text:  # ✅ Fix: Use pre-extracted text from frontend if available
            image = Image.open(io.BytesIO(file.read()))
            extracted_text = pytesseract.image_to_string(image, lang=language).strip()

        if not extracted_text:
            return jsonify({"error": "No text detected"}), 400

        filename = os.path.splitext(file.filename)[0] + ".xlsx"
        excel_path = save_text_as_excel(extracted_text, filename)

        return jsonify({"download_url": f"/download/{filename}"})

    except Exception as e:
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500


# Function to convert PDF to Images
def convert_pdf_to_images(pdf_bytes, filename_base):
    images = convert_from_bytes(pdf_bytes)  # Convert PDF to images
    image_paths = []

    # Save images and store paths
    for idx, image in enumerate(images):
        img_filename = f"{filename_base}_page_{idx + 1}.png"
        img_path = os.path.join(OUTPUT_DIR, img_filename)
        image.save(img_path, "PNG")
        image_paths.append(img_path)

    return image_paths

# Route for PDF to Image Conversion
@app.route("/pdf-to-images", methods=["POST"])
def convert_pdf_images():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]

    try:
        pdf_bytes = file.read()

        # Process PDF to Images
        filename_base = secure_filename(file.filename.rsplit(".", 1)[0])

        # Ensure filename_base is not empty
        if not filename_base:
            filename_base = "pdf_page"  # Default filename for PDF pages

        # Convert PDF to Images
        image_paths = convert_pdf_to_images(pdf_bytes, filename_base)

        # Generate download URLs
        download_urls = [f"/download/{os.path.basename(img_path)}" for img_path in image_paths]

        return jsonify({"download_urls": download_urls})  # Return a list of URLs

    except Exception as e:
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500


UPLOAD_FOLDER = "uploads"  # Ensure this matches your actual upload folder
# ✅ Route for Downloading Files
@app.route("/download/<filename>")
def download_file(filename):
    file_path = os.path.join(OUTPUT_DIR, filename)

    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404

    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/features')
def features():
    return render_template('features.html')
@app.route('/faq')
def faq():
    return render_template('faq.html')

@app.route('/pricing')
def pricing():
    return render_template('pricing.html')

@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/term')
def term():
    return render_template('term.html')

@app.route('/api', endpoint='api_page')
def api():
    return render_template('api.html')



port = int(os.environ.get("PORT", 10000))  # fallback doesn't matter, Render overrides it

if __name__ == "__main__":
    print(f"Running on port {port}")  # Optional debug print
    app.run(host="0.0.0.0", port=port)

